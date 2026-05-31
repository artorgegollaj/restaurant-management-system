# -*- coding: utf-8 -*-
"""Generate the presentation-prep Word document on the Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ORANGE = RGBColor(0xE5, 0x5A, 0x00)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
RED = RGBColor(0xC0, 0x00, 0x00)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, level=1, color=ORANGE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def para(text="", bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = DARK
    return p

# ---------------- Title page ----------------
t = doc.add_heading("Restaurant Management System (AURA)", level=0)
for run in t.runs:
    run.font.color.rgb = ORANGE
sub = para("Presentation Prep — Architecture, Tokens, Roles & Per-Member Guide", italic=True, size=13)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = para("Spring Boot 3 (Java 17) + React (Vite + MobX) + MSSQL + JWT/Refresh tokens", size=10)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------- Critical warnings ----------------
heading("Critical Warnings (read first)", level=1, color=RED)
para("Login password", bold=True)
bullet("The app actually seeds admin / admin123 (DataSeeder.java -> encoder.encode(\"admin123\"), and Login.jsx pre-fills it). The demo-script.md and qa-prep.md say Admin@12345 which is WRONG. Use admin123 in the demo and fix those two docs.")
para("Backend won't compile right now", bold=True)
bullet("The unfinished delivery feature left DashboardController depending on a DeliveryRepository class that doesn't exist, and DeliveryControllerTest.java is in the wrong folder. Until that is finished or stubbed, mvn spring-boot:run fails and the frontend has no working API. The React UI runs fine on its own.")

# ---------------- Part 1 ----------------
heading("Part 1 - The Whole Project in One Picture", level=1)
para("Three tiers:", bold=True)
code("React + Vite + MobX  --HTTP/JSON, Authorization: Bearer-->  Spring Boot 3 (Java 17)  --JPA/Hibernate-->  MSSQL (RestaurantDB)\n   (port 5173/5174)                                                  (port 8080)")
bullet("controllers (REST) -> repositories (Spring Data JPA) -> MSSQL. Security via Spring Security + a JWT filter. Tables auto-created by Hibernate (ddl-auto=update).", bold_prefix="Backend: ")
bullet("pages + a generic CrudPage driven by crudConfigs.jsx (one config per entity), an AuthStore (MobX) holding auth state in memory, and an axios instance (http.js) that attaches the token and auto-refreshes on 401.", bold_prefix="Frontend: ")
bullet("User, Role, RefreshToken, MenuCategory, MenuItem, RestaurantTable, OrderEntity, OrderItem, Payment, Reservation, Review, Staff, Ingredient (+ the in-progress Delivery).", bold_prefix="Entities (13): ")

# ---------------- Part 2 ----------------
heading("Part 2 - Tokens & Roles (the part the professor WILL drill)", level=1)
para("This is the most important section. Everyone must be able to say this.", italic=True)

heading("Access token", level=2)
bullet("a JWT, signed with HMAC-SHA using the secret in application.properties (app.jwt.secret, base64).", bold_prefix="What: ")
bullet("JwtService.generateAccessToken(username) - it puts only the username (subject), issuedAt, and expiration inside.", bold_prefix="Made by: ")
bullet("app.jwt.access-expiration-ms = 900000 = 15 minutes.", bold_prefix="Lifetime: ")
bullet("only in memory on the frontend, in AuthStore.accessToken (MobX). NOT in localStorage, sessionStorage, or cookies -> gone on a hard refresh (F5). This is the anti-XSS design.", bold_prefix="Stored where: ")
bullet("the axios request interceptor in http.js adds Authorization: Bearer <accessToken> to every API call.", bold_prefix="Sent how: ")
bullet("JwtAuthFilter reads the header -> JwtService.isValid() verifies signature + expiry -> loads the user from DB -> fills the SecurityContext.", bold_prefix="Checked how: ")

heading("Refresh token", level=2)
bullet("an opaque random string (two UUIDs joined), NOT a JWT.", bold_prefix="What: ")
bullet("app.jwt.refresh-expiration-ms = 604800000 = 7 days.", bold_prefix="Lifetime: ")
para("Stored where - TWO places:", bold=True)
bullet("In the database (MSSQL), table refresh_tokens - columns: token, user_id (FK), expiryDate, revoked. This DB copy is what lets us revoke it.")
bullet("In memory on the frontend, AuthStore.refreshToken - again not in localStorage/cookies.")
bullet("when the access token expires, an API call returns 401; the axios response interceptor in http.js calls authStore.refresh() -> POST /api/auth/refresh -> backend looks it up in DB, checks valid (not revoked, not expired), revokes the old one and issues a new pair (rotation), retries the original request.", bold_prefix="Used how: ")
bullet("on logout -> POST /api/auth/logout sets revoked = true in the DB row. There's also logout-all (revoke every token for a user).", bold_prefix="Revoked how: ")

heading("Roles", level=2)
bullet("in the DB - table roles + the join table user_roles (a @ManyToMany between User and Role). Seeded roles: ADMIN, MANAGER, WAITER, USER.", bold_prefix="Stored where: ")
bullet("roles are NOT inside the JWT. On every request, JwtAuthFilter -> CustomUserDetailsService.loadUserByUsername() reloads the user and roles from the DB and maps them to ROLE_ADMIN, etc. Trade-off: always fresh roles, but one DB read per request.", bold_prefix="Important: ")
bullet("the roles also come back to the frontend at login and sit in AuthStore.roles, but that's only for UI routing. Real security is enforced on the backend with @EnableMethodSecurity + @PreAuthorize on controller methods.", bold_prefix="Enforcement: ")

heading("The exact answer to: Ku e keni ruajtur refresh tokenin? (memorize)", level=2)
p = doc.add_paragraph()
r = p.add_run("Refresh token-i ruhet ne dy vende: (1) ne databazen MSSQL, ne tabelen refresh_tokens, me nje expiryDate dhe nje flag revoked - kjo na lejon ta anulojme menjehere ne logout; dhe (2) ne frontend vetem ne memorie, ne MobX AuthStore - JO ne localStorage ose cookies, per te shmangur sulmet XSS. Access token-i ruhet vetem ne memorie ne frontend dhe zgjat 15 minuta; refresh token-i zgjat 7 dite dhe eshte nje string i rastesishem (UUID), jo JWT.")
r.italic = True
para("If asked \"a i mbani rolet brenda JWT-se?\" -> \"Jo. JWT-ja mban vetem username-in. Rolet i lexojme nga databaza ne cdo kerkese permes CustomUserDetailsService.\"", italic=True)

# ---------------- Part 3 ----------------
heading("Part 3 - What Each Member Did + Personal Q&A", level=1)
para("Each person owned 6 cards plus a slice of the delivery feature.", italic=True)

members = [
    ("1) Yll Nuradini (you)", [
        "Orders backend - OrderController.create() accepts a whole order with nested items in one request and auto-calculates the total (price x quantity summed).",
        "Loading spinners / skeleton states in tables - Spinner.jsx, TableSkeleton.jsx.",
        "Seed data - DataSeeder.java: roles, admin user, 4 categories, 5 menu items, 10 tables.",
        "Auth integration tests - AuthControllerTest.java (register -> login -> refresh).",
        "README - architecture + entity diagram + setup.",
        "Demo script - team-tasks/demo-script.md.",
    ], [
        "How is the order total calculated? -> backend loops items, total += menuItem.price x quantity, ignores any client-sent total.",
        "Where does the seed admin come from? -> DataSeeder (a CommandLineRunner) runs on startup, password BCrypt-hashed.",
        "Live task: Create an order. -> Orders -> New -> pick table, add items -> Save -> show total.",
    ]),
    ("2) Artor Gegollaj", [
        "Global exception handler - GlobalExceptionHandler.java (@ControllerAdvice) turns errors into clean JSON; handles validation, bad credentials (401), access denied (403).",
        "@PreAuthorize role rules on controllers (the actual authorization).",
        "Status badges with color coding - StatusBadge.jsx.",
        "database/schema.sql - full DDL + indexes + constraints.",
        "Roles & permissions table (docs - who can do what).",
        "Backup demo video.",
    ], [
        "What happens when a WAITER tries to delete a menu item? -> @PreAuthorize(\"hasRole('ADMIN')\") blocks it -> AccessDeniedException -> handler returns 403.",
        "How do you handle errors globally? -> @ControllerAdvice + @ExceptionHandler.",
        "Live task: log in as a non-admin and show the 403.",
    ]),
    ("3) Yll Kadrijaj", [
        "PATCH /api/orders/{id}/status with a state machine - valid statuses + allowed transitions (PENDING->IN_PROGRESS->DELIVERED->PAID, or CANCELLED); invalid transition -> 409 Conflict.",
        "Image upload component - ImageUpload.jsx (multipart).",
        "Form validation messages - useFormValidation.js + FormField.jsx (Bootstrap invalid-feedback).",
        "Swagger / OpenAPI - OpenApiConfig.java, UI at /swagger-ui.html.",
        "Order tests - OrderControllerTest.java.",
        "Q&A prep doc.",
    ], [
        "Why can't an order go straight from PENDING to PAID? -> ALLOWED_TRANSITIONS only permits PENDING->IN_PROGRESS or CANCELLED; anything else throws 409.",
        "Show me the API docs. -> open http://localhost:8080/swagger-ui.html.",
        "Live task: change an order's status through allowed steps and show a rejected jump.",
    ]),
    ("4) Auron Blakaj", [
        "Bean Validation annotations (@NotBlank, @Email, @Size, @DecimalMin) on DTOs/entities - e.g. register request, MenuItem.name/price.",
        "Star rating component for Reviews - StarRating.jsx.",
        "DB CHECK constraints on enum columns (Status, Shift, Method).",
        "Postman collection for all endpoints.",
        "Demo-data reset script - demo-reset.sh.",
        "Final CRUD walkthrough test (in demo order).",
    ], [
        "What if I submit a menu item with an empty name? -> @NotBlank(message=\"Emertimi i detyrueshem\") -> 400 with that message (via the global handler).",
        "How do you validate input? -> Bean Validation at the DTO layer + custom checks in controllers (valid statuses, allowed transitions).",
        "Live task: add a Review and show the clickable 1-5 stars.",
    ]),
    ("5) Drion Gegollaj", [
        "Dashboard backend - DashboardController: daily-sales?from=&to= (date filter), top-products, summary, sales-report. Restricted to ADMIN/MANAGER.",
        "Logout revokes refresh token in DB - AuthController.logout().",
        "Dashboard charts - Chart.js (7-day sales + top products).",
        "Toast notifications - react-toastify (replacing alert()).",
        "Cross-browser test checklist.",
        "Slides outline - presentation-outline.md.",
    ], [
        "How is daily sales computed? -> filters PAID orders in the date range, groups by day, sums totals (fills empty days with 0).",
        "What does logout actually do? -> finds the refresh token in DB and sets revoked=true; query SELECT * FROM refresh_tokens WHERE revoked=1 to prove it.",
        "Live task: open Dashboard, filter by date range.",
    ]),
]

for name, built, qa in members:
    heading(name, level=2)
    para("Built:", bold=True)
    for b in built:
        bullet(b)
    para("Be ready to explain / likely questions:", bold=True)
    for q in qa:
        bullet(q)

heading("Delivery feature (shared, UNFINISHED)", level=2)
bullet("Split across all five (06-delivery-feature.md). The Delivery entity + DeliveryRepository aren't fully in place, which is why the backend currently doesn't compile. Either finish it or temporarily remove the delivery-stats / DeliveryRepository references from DashboardController before the demo. Don't let the professor catch the app not starting - decide this in advance.")

# ---------------- Part 4 ----------------
heading("Part 4 - Live CRUD (if asked: add/update/delete a menu item)", level=1)
para("Who is allowed (backend @PreAuthorize):", bold=True)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 2"
hdr = table.rows[0].cells
hdr[0].text = "Action"; hdr[1].text = "Endpoint"; hdr[2].text = "Allowed roles"
rows = [
    ("List / view", "GET /api/menu-items", "any logged-in user"),
    ("Create", "POST /api/menu-items", "ADMIN, MANAGER"),
    ("Update", "PUT /api/menu-items/{id}", "ADMIN, MANAGER"),
    ("Delete", "DELETE /api/menu-items/{id}", "ADMIN only"),
]
for a, b, c in rows:
    cells = table.add_row().cells
    cells[0].text = a; cells[1].text = b; cells[2].text = c
para("Since you log in as admin (ROLE_ADMIN), you can do all three.")

para("Add a menu item (live):", bold=True)
bullet("Sidebar -> Menu Items -> Add/New.")
bullet("Fill Name (required - empty name -> red \"Emertimi i detyrueshem\"), Description, Price (>= 0), Image URL, Available checkbox, and pick a Category from the dropdown (loaded live from /menu-categories).")
bullet("Save -> POST /api/menu-items. Open DevTools -> Network to show the request + Authorization: Bearer header.")
para("Update:", bold=True)
bullet("click a row's edit -> change price -> Save -> PUT /api/menu-items/{id} (backend copies fields onto the existing row and re-saves).")
para("Delete:", bold=True)
bullet("click delete -> DELETE /api/menu-items/{id} (ADMIN-only).")
para("Talking point if they push:", bold=True)
bullet("Categories have no role restriction (any logged-in user), but menu items require ADMIN/MANAGER to write and ADMIN to delete - enforced by @PreAuthorize on each method, not by the frontend.")

# ---------------- How to run ----------------
heading("Appendix - How to Run", level=1)
bullet("DB: SQL Server on localhost:1433, database RestaurantDB, login restaurant_app / Restaurant@2026.", bold_prefix="1. ")
bullet("Backend: cd backend && .\\mvnw.cmd spring-boot:run  (starts on :8080, seeds admin/admin123).", bold_prefix="2. ")
bullet("Frontend: cd frontend && npm install && npm run dev  (starts on :5173, or :5174 if busy).", bold_prefix="3. ")
bullet("Open the frontend URL in Chrome and log in with admin / admin123.", bold_prefix="4. ")

out = r"C:\Users\DELL\Desktop\Presentation-Prep.docx"
doc.save(out)
print("SAVED:", out)
