# -*- coding: utf-8 -*-
"""Gjeneron dokumentin e pergatitjes per prezantim (shqip) ne Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ORANGE = RGBColor(0xE5, 0x5A, 0x00)
DARK = RGBColor(0x1a, 0x1a, 0x1a)
RED = RGBColor(0xC0, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, level=1, color=ORANGE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def para(text="", bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = DARK
    return p

def qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run("P: " + q); r.bold = True; r.font.color.rgb = BLUE
    pa = doc.add_paragraph()
    ra = pa.add_run("PÃ«rgjigje: ".replace("Ã«","ë") + a)

# ---------------- Kopertina ----------------
t = doc.add_heading("Restaurant Management System (AURA)", level=0)
for run in t.runs: run.font.color.rgb = ORANGE
sub = para("PÃ«rgatitje pÃ«r Prezantim — Arkitektura, Tokens, Roles & Udhëzues për secilin anëtar".replace("Ã«","ë"), italic=True, size=13)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = para("Spring Boot 3 (Java 17) + React (Vite + MobX) + MSSQL + JWT/Refresh tokens", size=10)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------- Paralajmerime ----------------
heading("Paralajmërime kritike (lexoni të parat)", level=1, color=RED)
para("Password-i i login-it", bold=True)
bullet("Aplikacioni e bën seed admin / admin123 (DataSeeder.java -> encoder.encode(\"admin123\"), dhe Login.jsx e ka të para-mbushur). Dokumentet demo-script.md dhe qa-prep.md thonë Admin@12345 që është GABIM. Përdorni admin123 në demo dhe rregulloni ato dy dokumente.")
para("Backend-i nuk kompajllon për momentin", bold=True)
bullet("Feature-i delivery i papërfunduar e la DashboardController të varur nga një klasë DeliveryRepository që nuk ekziston, dhe DeliveryControllerTest.java është në folderin e gabuar. Derisa kjo të përfundojë ose të bëhet stub, mvn spring-boot:run dështon dhe frontend-i nuk ka API funksionale. UI-ja React punon vetë.")

# ---------------- Pjesa 1 ----------------
heading("Pjesa 1 - I gjithë projekti në një pamje", level=1)
para("Tri shtresa (three tiers):", bold=True)
code("React + Vite + MobX  --HTTP/JSON, Authorization: Bearer-->  Spring Boot 3 (Java 17)  --JPA/Hibernate-->  MSSQL (RestaurantDB)\n   (port 5173/5174)                                                  (port 8080)")
bullet("controllers (REST) -> repositories (Spring Data JPA) -> MSSQL. Siguria me Spring Security + një JWT filter. Tabelat krijohen automatikisht nga Hibernate (ddl-auto=update).", bold_prefix="Backend: ")
bullet("pages + një CrudPage gjenerik i drejtuar nga crudConfigs.jsx (një config për çdo entity), një AuthStore (MobX) që mban gjendjen e auth në memorie, dhe një instance axios (http.js) që bashkëngjit token-in dhe bën auto-refresh në 401.", bold_prefix="Frontend: ")
bullet("User, Role, RefreshToken, MenuCategory, MenuItem, RestaurantTable, OrderEntity, OrderItem, Payment, Reservation, Review, Staff, Ingredient (+ Delivery që është në punë).", bold_prefix="Entities (13): ")

# ---------------- Pjesa 2 ----------------
heading("Pjesa 2 - Tokens & Roles (pjesa që profesori DO ta pyesë)", level=1)
para("Kjo është pjesa më e rëndësishme. Të gjithë duhet ta dini këtë.", italic=True)

heading("Access token", level=2)
bullet("një JWT, i nënshkruar me HMAC-SHA duke përdorur secret-in te application.properties (app.jwt.secret, base64).", bold_prefix="Çfarë është: ")
bullet("JwtService.generateAccessToken(username) - vendos vetëm username-in (subject), issuedAt dhe expiration brenda.", bold_prefix="Krijohet nga: ")
bullet("app.jwt.access-expiration-ms = 900000 = 15 minuta.", bold_prefix="Jetëgjatësia: ")
bullet("vetëm në memorie në frontend, te AuthStore.accessToken (MobX). JO në localStorage, sessionStorage ose cookies -> humbet me një hard refresh (F5). Ky është dizajni anti-XSS.", bold_prefix="Ku ruhet: ")
bullet("request interceptor i axios te http.js shton Authorization: Bearer <accessToken> në çdo API call.", bold_prefix="Si dërgohet: ")
bullet("JwtAuthFilter lexon header-in -> JwtService.isValid() verifikon signature + expiry -> ngarkon user-in nga DB -> mbush SecurityContext.", bold_prefix="Si verifikohet: ")

heading("Refresh token", level=2)
bullet("një string i rastësishëm (dy UUID të bashkuar), JO një JWT.", bold_prefix="Çfarë është: ")
bullet("app.jwt.refresh-expiration-ms = 604800000 = 7 ditë.", bold_prefix="Jetëgjatësia: ")
para("Ku ruhet - NË DY VENDE:", bold=True)
bullet("Në databazë (MSSQL), tabela refresh_tokens - kolonat: token, user_id (FK), expiryDate, revoked. Kjo kopje në DB është ajo që na lejon ta bëjmë revoke.")
bullet("Në memorie në frontend, AuthStore.refreshToken - sërish jo në localStorage/cookies.")
bullet("kur access token-i skadon, një API call kthen 401; response interceptor i axios te http.js thërret authStore.refresh() -> POST /api/auth/refresh -> backend-i e gjen në DB, kontrollon nëse është valid (jo revoked, jo i skaduar), bën revoke të vjetrin dhe lëshon një çift të ri (rotation), riprovon kërkesën origjinale.", bold_prefix="Si përdoret: ")
bullet("në logout -> POST /api/auth/logout vendos revoked = true në rreshtin e DB. Ka edhe logout-all (revoke çdo token për një user).", bold_prefix="Si bëhet revoke: ")

heading("Roles", level=2)
bullet("në DB - tabela roles + tabela join user_roles (një @ManyToMany mes User dhe Role). Roles të seeded: ADMIN, MANAGER, WAITER, USER.", bold_prefix="Ku ruhen: ")
bullet("roles NUK janë brenda JWT-së. Në çdo kërkesë, JwtAuthFilter -> CustomUserDetailsService.loadUserByUsername() ringarkon user-in dhe roles nga DB dhe i mapon në ROLE_ADMIN, etj. Trade-off: roles gjithmonë të freskëta, por një DB read për çdo kërkesë.", bold_prefix="E rëndësishme: ")
bullet("roles kthehen edhe te frontend në login dhe rrinë te AuthStore.roles, por kjo është vetëm për UI routing. Siguria e vërtetë zbatohet në backend me @EnableMethodSecurity + @PreAuthorize mbi metodat e controller-ave.", bold_prefix="Zbatimi: ")

heading("Përgjigja e saktë për: Ku e keni ruajtur refresh tokenin? (mësojeni përmendësh)", level=2)
p = doc.add_paragraph()
r = p.add_run("Refresh token-i ruhet në dy vende: (1) në databazën MSSQL, në tabelën refresh_tokens, me një expiryDate dhe një flag revoked - kjo na lejon ta anulojmë menjëherë në logout; dhe (2) në frontend vetëm në memorie, në MobX AuthStore - JO në localStorage ose cookies, për të shmangur sulmet XSS. Access token-i ruhet vetëm në memorie në frontend dhe zgjat 15 minuta; refresh token-i zgjat 7 ditë dhe është një string i rastësishëm (UUID), jo JWT.")
r.italic = True
para("Nëse pyesin \"a i mbani rolet brenda JWT-së?\" -> \"Jo. JWT-ja mban vetëm username-in. Rolet i lexojmë nga databaza në çdo kërkesë përmes CustomUserDetailsService.\"", italic=True)

# ---------------- Pjesa 3 - in depth per member ----------------
heading("Pjesa 3 - Çfarë bëri secili (në detaje)", level=1)
para("Secili person zotëroi 6 karta plus një pjesë të feature-it delivery.", italic=True)

# Yll Nuradini
heading("1) Yll Nuradini", level=2)
para("Çfarë ndërtoi:", bold=True)
bullet("Orders backend - OrderController.create() pranon një porosi të tërë me nested items në një kërkesë dhe e llogarit totalin automatikisht. Backend-i nuk i beson totalit nga klienti: për çdo OrderItem merr menuItem nga DB, shumëzon price x quantity dhe i mbledh. Kontrollon edhe nëse items është bosh -> 400.")
bullet("Loading spinners / skeleton states në tabela - Spinner.jsx, TableSkeleton.jsx (UX gjatë ngarkimit).")
bullet("Seed data - DataSeeder.java (implementon CommandLineRunner): 4 roles, user-in admin (password BCrypt), 4 kategori, 5 menu items, 10 tavolina. Ka if-e që të mos dyfishojë.")
bullet("Auth integration tests - AuthControllerTest.java me @SpringBootTest + MockMvc (register -> login -> refresh), me H2 in-memory.")
bullet("README - arkitektura + entity diagram + setup.")
bullet("Demo script - team-tasks/demo-script.md.")

# Artor
heading("2) Artor Gegollaj", level=2)
para("Çfarë ndërtoi:", bold=True)
bullet("Global exception handler - GlobalExceptionHandler.java (@ControllerAdvice) i kthen gabimet në JSON të pastër {timestamp, status, message}; trajton validimin (400), bad credentials (401), access denied (403), dhe Exception e përgjithshme (500).")
bullet("@PreAuthorize role rules mbi controllers (autorizimi i vërtetë), i aktivizuar nga @EnableMethodSecurity.")
bullet("Status badges me color coding - StatusBadge.jsx.")
bullet("database/schema.sql - DDL i plotë + indexe + constraints.")
bullet("Tabela e roleve & lejeve (docs - kush mund të bëjë çfarë).")
bullet("Backup video i demo-s.")

# Kadrijaj
heading("3) Yll Kadrijaj", level=2)
para("Çfarë ndërtoi:", bold=True)
bullet("PATCH /api/orders/{id}/status me një state machine - VALID_STATUSES + ALLOWED_TRANSITIONS (PENDING->IN_PROGRESS->DELIVERED->PAID, ose CANCELLED); tranzicion i palejuar -> 409 Conflict.")
bullet("Image upload component - ImageUpload.jsx (multipart); backend e ruan te uploads/ dhe e shërben me WebConfig (/uploads/**).")
bullet("Form validation messages - useFormValidation.js + FormField.jsx (Bootstrap invalid-feedback).")
bullet("Swagger / OpenAPI - OpenApiConfig.java, UI te /swagger-ui.html (springdoc).")
bullet("Order tests - OrderControllerTest.java.")
bullet("Q&A prep doc.")

# Auron
heading("4) Auron Blakaj", level=2)
para("Çfarë ndërtoi:", bold=True)
bullet("Bean Validation annotations (@NotBlank, @Email, @Size, @DecimalMin) te DTO/entitetet - p.sh. register request, MenuItem.name/price - me @Valid te controller-i.")
bullet("Star rating component për Reviews - StarRating.jsx (1-5 yje të klikueshëm).")
bullet("DB CHECK constraints mbi kolonat enum (Status, Shift, Method) - mbrojtje në nivel DB.")
bullet("Postman collection E2E për të gjitha endpoint-et.")
bullet("Demo-data reset script - demo-reset.sh.")
bullet("Test final i të gjitha CRUD-eve sipas radhës së demo-s.")

# Drion
heading("5) Drion Gegollaj", level=2)
para("Çfarë ndërtoi:", bold=True)
bullet("Dashboard backend - DashboardController: daily-sales?from=&to= (date filter), top-products, summary, sales-report. I kufizuar te ADMIN/MANAGER.")
bullet("Logout revoke të refresh token në DB - AuthController.logout().")
bullet("Dashboard charts - Chart.js (shitjet 7 ditë + top produkte).")
bullet("Toast notifications - react-toastify (në vend të alert()).")
bullet("Checklist manual cross-browser test.")
bullet("Slides outline - presentation-outline.md.")

heading("Feature Delivery (i përbashkët, I PAPËRFUNDUAR)", level=2)
bullet("I ndarë mes të pesëve (06-delivery-feature.md). Entity Delivery + DeliveryRepository nuk janë plotësisht në vend, prandaj backend-i nuk kompajllon. Ose përfundojeni, ose hiqni përkohësisht referencat delivery-stats / DeliveryRepository nga DashboardController para demo-s.")

# ---------------- Pjesa 4 - Q&A per person (NEW) ----------------
heading("Pjesa 4 - Pyetje të mundshme nga profesori & Përgjigjet (për secilin)", level=1)
para("Pyetje specifike sipas pjesës që mbuloi secili, me përgjigje gati.", italic=True)

heading("Yll Nuradini - Orders, Seed data, Tests", level=2)
qa("Si llogaritet totali i porosisë?",
   "Në OrderController.create(), backend-i NUK i beson totalit që vjen nga klienti. Për çdo OrderItem e merr menuItem nga DB, shumëzon price x quantity dhe i mbledh të gjitha; pastaj e ruan te OrderEntity.total. Kjo parandalon manipulimin e çmimit nga frontend.")
qa("Çfarë ndodh nëse dërgohet një porosi pa artikuj?",
   "Kontrollohet nëse lista items është null ose bosh; nëse po, hedh ResponseStatusException me 400 BAD_REQUEST dhe mesazhin 'Porosia duhet të ketë të paktën një artikull'.")
qa("Si krijohen user-i admin dhe të dhënat fillestare?",
   "DataSeeder implementon CommandLineRunner, ekzekutohet automatikisht në startup. Krijon rolet (ADMIN, MANAGER, WAITER, USER), user-in admin me password të hash-uar me BCrypt, kategoritë, menu items dhe 10 tavolinat. Ka if-e (existsByUsername, count()==0) që të mos i dyfishojë.")
qa("Pse e quani entity-n OrderEntity dhe jo Order?",
   "Sepse 'Order' është fjalë e rezervuar në SQL/JPQL (ORDER BY). Duke e quajtur OrderEntity, query-t JPQL funksionojnë; tabela në DB mbetet 'orders'.")
qa("Si i testoni endpoint-et e autentikimit?",
   "Me AuthControllerTest: @SpringBootTest + MockMvc. Bën register -> login -> refresh dhe verifikon që login kthen accessToken dhe refreshToken jo-null. Testet përdorin H2 in-memory me profilin 'test', kështu nuk prekin DB-në reale.")

heading("Artor Gegollaj - Error handling & Security", level=2)
qa("Si i menaxhoni gabimet në mënyrë qendrore?",
   "Me GlobalExceptionHandler të anotuar @ControllerAdvice. Ka @ExceptionHandler për ResponseStatusException, MethodArgumentNotValidException (validim -> 400), BadCredentialsException (401), AccessDeniedException (403) dhe Exception e përgjithshme (500). Të gjitha kthejnë JSON {timestamp, status, message}.")
qa("Çfarë ndodh kur një WAITER provon të fshijë një menu item?",
   "Metoda delete ka @PreAuthorize(\"hasRole('ADMIN')\"). Spring Security hedh AccessDeniedException, e cila kapet nga handler-i dhe kthehet 403 me 'Nuk keni leje për këtë veprim'.")
qa("Ku përcaktohen rregullat e roleve?",
   "Me anotimin @PreAuthorize mbi metodat e controller-ave (p.sh. hasAnyRole('ADMIN','MANAGER')), i aktivizuar nga @EnableMethodSecurity te SecurityConfig.")
qa("Pse keni schema.sql nëse Hibernate i krijon tabelat vetë?",
   "ddl-auto=update është për zhvillim. schema.sql jep DDL eksplicit me indexe dhe constraints, e dokumenton strukturën dhe mund të përdoret për prod (ku do përdornim migrime si Flyway).")
qa("Si funksionon CORS në projekt?",
   "Te SecurityConfig, corsSource() lejon origin-at nga app.cors.origin (localhost:5173 dhe 5174), metodat GET/POST/PUT/DELETE/OPTIONS dhe të gjitha header-at. Kjo lejon frontend-in të thërrasë backend-in në port tjetër.")

heading("Yll Kadrijaj - Order status, Swagger, Validation UI", level=2)
qa("Si funksionon ndryshimi i statusit të porosisë?",
   "PATCH /api/orders/{id}/status me dy kontrolle: (1) VALID_STATUSES - a është status i njohur; (2) ALLOWED_TRANSITIONS - një Map që përcakton nga cili status mund të kalosh ku. Status i panjohur -> 400; tranzicion i palejuar -> 409 CONFLICT.")
qa("Pse një porosi s'mund të shkojë direkt nga PENDING në PAID?",
   "Sepse te ALLOWED_TRANSITIONS, PENDING lejon vetëm IN_PROGRESS ose CANCELLED. PAID arrihet vetëm pasi të kalojë DELIVERED. Kjo është një state machine që pasqyron rrjedhën reale.")
qa("Ku i shihni API docs dhe si i gjeneroni?",
   "Swagger UI te http://localhost:8080/swagger-ui.html, i konfiguruar te OpenApiConfig.java me springdoc. Endpoint-et lejohen te SecurityConfig (/v3/api-docs/**, /swagger-ui/**).")
qa("Si bëhet validimi i formave në frontend?",
   "Me hook-un custom useFormValidation.js + komponentin FormField.jsx, që shfaqin mesazhe gabimi me stilin Bootstrap invalid-feedback PARA se kërkesa të dërgohet.")
qa("Si funksionon image upload?",
   "ImageUpload.jsx dërgon file-in si multipart; backend-i e ruan te folderi uploads/ dhe e shërben statikisht përmes WebConfig (resource handler /uploads/**).")

heading("Auron Blakaj - Validation & Testing", level=2)
qa("Si validoni input-in në backend?",
   "Me Bean Validation: anotime si @NotBlank, @Email, @Size, @DecimalMin te DTO/entitetet, plus @Valid te parametri i controller-it. Nëse validimi dështon, GlobalExceptionHandler kthen 400 me mesazhin e fushës.")
qa("Çfarë ndodh nëse dërgoj një menu item me emër bosh?",
   "@NotBlank(message=\"Emertimi i detyrueshem\") dështon -> përgjigjja 400 me atë mesazh (e formatuar nga GlobalExceptionHandler).")
qa("Çfarë janë CHECK constraints dhe pse i shtuat?",
   "Janë kufizime në nivel DB që lejojnë vetëm vlera enum të vlefshme (p.sh. status, shift, method). Janë një shtresë e dytë mbrojtjeje përveç validimit në aplikacion - edhe nëse dikush shkruan direkt në DB, vlerat e gabuara refuzohen.")
qa("Si i keni testuar të gjitha endpoint-et?",
   "Me një Postman collection që mbulon CRUD-et e të gjitha entiteteve, plus një demo-reset script për të rivendosur të dhënat para prezantimit.")
qa("Çfarë është star rating component?",
   "StarRating.jsx - një komponent React me 1-5 yje të klikueshëm, i përdorur te Reviews për të dhënë/treguar vlerësimin.")

heading("Drion Gegollaj - Dashboard, Logout, Frontend UX", level=2)
qa("Si llogariten shitjet ditore?",
   "GET /api/dashboard/daily-sales?from=&to=. Filtron porositë me status PAID brenda intervalit, i grupon sipas datës, mbledh totalet dhe i mbush ditët bosh me 0. Endpoint-i është i kufizuar te ADMIN/MANAGER.")
qa("Çfarë bën logout-i konkretisht?",
   "AuthController.logout() e gjen refresh token-in në DB dhe vendos revoked=true. Pas kësaj ai token s'pranohet më në /refresh. Provohet me SELECT * FROM refresh_tokens WHERE revoked=1.")
qa("Si i llogaritni top produktet?",
   "Endpoint-i top-products agregon të gjitha OrderItem-at sipas emrit të menu item, mbledh sasinë totale dhe të ardhurat, i rendit zbritës sipas sasisë dhe kthen top N (default 5).")
qa("Pse përdorni react-toastify në vend të alert()?",
   "Toast-et janë njoftime jo-bllokuese dhe më profesionale; alert() i browser-it bllokon UI-në dhe e ndërpret rrjedhën e eventeve.")
qa("Si i shfaqni grafikët në Dashboard?",
   "Me Chart.js, duke marrë të dhëna nga endpoint-et e dashboard-it (daily-sales dhe top-products) dhe duke i vizatuar si grafik shitjesh 7-ditor dhe top produkte.")

# ---------------- Pjesa 5 - Live CRUD ----------------
heading("Pjesa 5 - Live CRUD (nëse kërkojnë: shto/ndrysho/fshi një menu item)", level=1)
para("Ku ta gjeni në kod:", bold=True)
bullet("Backend: backend/src/main/java/com/ubt/restaurant/controller/MenuItemController.java -> create() (POST), update() (PUT), delete() (DELETE).")
bullet("Entity (fushat): entity/MenuItem.java (name, description, price, image, category, available).")
bullet("Frontend (forma): crudConfigs.jsx -> blloku menuItems (këtu shton një fushë të re 'on the fly').")
bullet("Frontend (UI gjenerik): components/CrudPage.jsx -> tabela + butonat Add/Edit/Delete.")

para("Kush lejohet (backend @PreAuthorize):", bold=True)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 2"
hdr = table.rows[0].cells
hdr[0].text = "Veprimi"; hdr[1].text = "Endpoint"; hdr[2].text = "Roles të lejuara"
rows = [
    ("Listo / shiko", "GET /api/menu-items", "çdo user i loguar"),
    ("Krijo", "POST /api/menu-items", "ADMIN, MANAGER"),
    ("Ndrysho", "PUT /api/menu-items/{id}", "ADMIN, MANAGER"),
    ("Fshi", "DELETE /api/menu-items/{id}", "vetëm ADMIN"),
]
for a, b, c in rows:
    cells = table.add_row().cells
    cells[0].text = a; cells[1].text = b; cells[2].text = c
para("Meqë loginoheni si admin (ROLE_ADMIN), i bëni të treja.")

para("Shto një menu item (live):", bold=True)
bullet("Menu anësore -> Menu Items -> Add/New.")
bullet("Plotëso Name (i detyrueshëm - bosh -> 'Emertimi i detyrueshem'), Description, Price (>= 0), Image URL, checkbox Available, dhe zgjidh Category nga dropdown (ngarkohet live nga /menu-categories).")
bullet("Save -> POST /api/menu-items. Hap DevTools -> Network për të treguar kërkesën + header-in Authorization: Bearer.")
para("Ndrysho:", bold=True)
bullet("edito rreshtin -> ndrysho çmimin -> Save -> PUT /api/menu-items/{id} (backend kopjon fushat te rreshti ekzistues dhe e ruan).")
para("Fshi:", bold=True)
bullet("butoni delete -> DELETE /api/menu-items/{id} (vetëm ADMIN).")

# ---------------- Shtojce ----------------
heading("Shtojcë - Si të ekzekutohet", level=1)
bullet("DB: SQL Server në localhost:1433, databaza RestaurantDB, login restaurant_app / Restaurant@2026.", bold_prefix="1. ")
bullet("Backend: cd backend && .\\mvnw.cmd spring-boot:run  (niset në :8080, seed admin/admin123).", bold_prefix="2. ")
bullet("Frontend: cd frontend && npm install && npm run dev  (niset në :5173, ose :5174 nëse është i zënë).", bold_prefix="3. ")
bullet("Hap URL-në e frontend në Chrome dhe logohu me admin / admin123.", bold_prefix="4. ")

out = r"C:\Users\DELL\Desktop\Pergatitje-Prezantimi-SQ.docx"
doc.save(out)
print("SAVED:", out)
